#!/usr/bin/env python3
"""Apply a list of edit operations to a base .docx, as a redline or a clean edit.

The deterministic *apply* helper for the edit-docx skill. The agent does
the *understand* step (reading the change source and deriving the operation list);
this script applies them to the base OOXML in place, matching each operation to a
paragraph by an anchor snippet. Two output modes:

- **tracked** (default) — record each change as Word tracked changes
  (``<w:ins>`` / ``<w:del>``) rather than a clean overwrite, so a reviewer opening
  the result in Word sees every insertion and deletion and can accept or reject each.
  A ``replace`` is word-diffed against the original paragraph, so only the words that
  actually changed are struck and inserted; the untouched wording stays live text.
  Pass ``--whole-paragraph`` to fall back to striking the entire paragraph instead.
- **clean** (``--clean``) — apply each change directly, producing a *final*
  document with the edits baked in and no tracked-change markup. Use this once the
  edits are settled (e.g. after the user has accepted/rejected a redline's
  suggestions) to hand back the final contract.

Either way styles, numbering, and fonts survive.

Usage::

    python apply_operations.py --original in.docx --ops ops.json --output out.docx \
        [--author "Name"] [--clean] [--whole-paragraph]

``ops.json`` shape::

    {"operations": [
        {"op": "replace",      "anchor_text": "...", "new_text": "..."},
        {"op": "delete",       "anchor_text": "..."},
        {"op": "insert_after", "anchor_text": "...", "new_text": "..."},
        {"op": "append",                             "new_text": "..."}
    ]}

Prints a JSON report to stdout: per-operation ``matched`` flag plus an
``unmatched`` count, so the caller can surface operations that found no anchor
(a common sign the change referenced text the base doesn't contain). Tracked
``replace`` operations also report the ``granularity`` actually used (``word``,
or ``paragraph`` when the paragraph's content forced the whole-paragraph fallback).
"""

from __future__ import annotations

import argparse
import copy
import datetime
import difflib
import itertools
import json
import re
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


# Inline markdown emphasis/code markers, and leading block markers (heading,
# ordered/unordered list, blockquote). The text the agent reads is markdown
# (the base is extracted to markdown at ingest), but python-docx paragraph text
# is raw — no `**`, no `3. ` list prefix. Stripping both lets an anchor copied
# verbatim from the markdown match the raw paragraph on the first try.
_MD_INLINE = re.compile(r"[*_`~]")
_MD_LEADING = re.compile(r"^\s*(?:#{1,6}\s+|\d+[.)]\s+|[-*+]\s+|>\s+)")

# Diff tokens for the word-level redline. Alternatives are ordered, and the
# concatenation of all matches reproduces the input exactly — token boundaries
# double as character offsets into the paragraph text. CJK is matched one
# character per token (no spaces to split on), so a Chinese contract redlines at
# the same resolution as an English one rather than as one giant token.
_TOKEN = re.compile(
    r"\s+"
    r"|[　-〿㐀-䶿一-鿿豈-﫿＀-￯]"
    r"|[0-9]+(?:[.,][0-9]+)*"
    r"|[^\W\d_]+(?:['’-][^\W\d_]+)*"
    r"|.",
    re.UNICODE,
)

# An unchanged stretch shorter than this (non-whitespace chars) that sits between
# two changes is absorbed into them, so "30 days" -> "45 days" reads as one edit
# instead of two strikeouts around a stranded " day".
_MIN_EQUAL_CHARS = 4

# Monotonic revision ids. Word requires every <w:ins>/<w:del> to carry a unique
# w:id within the document.
_rev_ids = itertools.count(1)


def _norm(text: str) -> str:
    """Normalize for anchor matching: strip markdown, collapse whitespace, lowercase.

    Tolerates reflow (whitespace), case, and the markdown-vs-raw gap between the
    text the agent reads and the raw paragraph text this script matches against.
    """
    text = _MD_LEADING.sub("", text or "")
    text = _MD_INLINE.sub("", text)
    return re.sub(r"\s+", " ", text).strip().lower()


def _find_paragraph(doc, anchor_text: str):
    """Return the first paragraph whose normalized text contains the anchor, or None."""
    needle = _norm(anchor_text)
    if not needle:
        return None
    for para in doc.paragraphs:
        if needle in _norm(para.text):
            return para
    return None


def _revision(tag: str, author: str, date: str):
    """Build an empty ``<w:ins>`` or ``<w:del>`` wrapper with a unique id/author/date."""
    el = OxmlElement(tag)
    el.set(qn("w:id"), str(next(_rev_ids)))
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), date)
    return el


def _make_text_run(new_text: str, rpr=None):
    """Build a ``<w:r>`` carrying *new_text*, optionally cloning run properties."""
    run = OxmlElement("w:r")
    if rpr is not None:
        run.append(copy.deepcopy(rpr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = new_text
    run.append(t)
    return run


def _make_del_run(text: str, rpr=None):
    """Build a ``<w:r>`` of struck-through text (``<w:delText>``) for a ``<w:del>``."""
    run = OxmlElement("w:r")
    if rpr is not None:
        run.append(copy.deepcopy(rpr))
    t = OxmlElement("w:delText")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    return run


def _first_run_rpr(para: Paragraph):
    """Return a copy target of the first run's ``<w:rPr>`` (or None), for format reuse."""
    if para.runs:
        return para.runs[0]._element.find(qn("w:rPr"))
    return None


def _mark_run_deleted(run, wrapper) -> None:
    """Move *run* into a ``<w:del>`` wrapper and convert its text to ``<w:delText>``.

    Deleted text lives in ``<w:delText>`` rather than ``<w:t>`` so Word renders it
    as struck-through tracked deletion instead of live text.
    """
    r = run._element
    for t in r.findall(qn("w:t")):
        t.tag = qn("w:delText")
    parent = r.getparent()
    parent.insert(parent.index(r), wrapper)
    wrapper.append(r)


def _mark_para_mark(para: Paragraph, tag: str, author: str, date: str) -> None:
    """Mark the paragraph mark itself as inserted/deleted (``<w:pPr><w:rPr><tag/>``).

    Needed so a whole-paragraph insert or delete tracks cleanly — accepting the
    revision also adds/removes the paragraph break, not just the visible text.
    """
    p = para._p
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        pPr.append(rPr)
    mark = OxmlElement(tag)
    mark.set(qn("w:id"), str(next(_rev_ids)))
    mark.set(qn("w:author"), author)
    mark.set(qn("w:date"), date)
    rPr.insert(0, mark)


def _run_spans(para: Paragraph):
    """Map the paragraph's runs to ``(element, rPr, text, start, end)`` char spans.

    Returns ``None`` when the paragraph holds content this script can't slice by
    character offset — a tab, a line break, a field, a drawing, or text living
    inside a hyperlink rather than a direct ``<w:r>`` child. Those paragraphs fall
    back to a whole-paragraph redline instead of a word diff, so nothing is dropped.
    """
    spans = []
    pos = 0
    for run in para.runs:
        el = run._element
        for child in el:
            if child.tag not in (qn("w:rPr"), qn("w:t")):
                return None
        text = run.text
        spans.append((el, el.find(qn("w:rPr")), text, pos, pos + len(text)))
        pos += len(text)
    if not spans or "".join(s[2] for s in spans) != para.text:
        return None  # empty, or text hiding in a hyperlink
    return spans


def _slice_spans(spans, start: int, end: int):
    """Return ``(rPr, text)`` pieces covering ``[start, end)``, preserving formatting."""
    pieces = []
    for _el, rpr, text, s, e in spans:
        if e <= start or s >= end:
            continue
        piece = text[max(start, s) - s : min(end, e) - s]
        if piece:
            pieces.append((rpr, piece))
    return pieces


def _coalesce(opcodes, old_tokens):
    """Merge changed opcode runs, absorbing the trivial equal blocks between them."""
    ops = list(opcodes)
    absorbed = set()
    for k, (tag, i1, i2, _j1, _j2) in enumerate(ops):
        if tag != "equal" or k == 0 or k == len(ops) - 1:
            continue  # a leading/trailing equal block has only one changed neighbour
        kept = re.sub(r"\s+", "", "".join(old_tokens[i1:i2]))
        if len(kept) < _MIN_EQUAL_CHARS:
            absorbed.add(k)

    merged, k = [], 0
    while k < len(ops):
        if ops[k][0] == "equal" and k not in absorbed:
            merged.append(ops[k])
            k += 1
            continue
        start = k
        while k < len(ops) and (ops[k][0] != "equal" or k in absorbed):
            k += 1
        merged.append(("change", ops[start][1], ops[k - 1][2], ops[start][3], ops[k - 1][4]))
    return merged


def _diff_segments(old_text: str, new_text: str):
    """Word-diff two strings into ``(kind, old_start, old_end, inserted_text)`` segments.

    ``old_start``/``old_end`` are character offsets into *old_text*, so each segment
    maps straight back onto the paragraph's runs.
    """
    old_tokens = _TOKEN.findall(old_text)
    new_tokens = _TOKEN.findall(new_text)

    offsets, pos = [], 0
    for token in old_tokens:
        offsets.append(pos)
        pos += len(token)
    offsets.append(pos)

    matcher = difflib.SequenceMatcher(None, old_tokens, new_tokens, autojunk=False)
    segments = []
    for tag, i1, i2, j1, j2 in _coalesce(matcher.get_opcodes(), old_tokens):
        old_slice = "".join(old_tokens[i1:i2])
        inserted = "".join(new_tokens[j1:j2])
        kind = "equal" if (tag == "equal" or old_slice == inserted) else "change"
        segments.append((kind, offsets[i1], offsets[i2], "" if kind == "equal" else inserted))
    return segments


def _replace_word_tracked(para: Paragraph, new_text: str, author: str, date: str) -> bool:
    """Redline only the words that changed; return False if the paragraph can't be diffed."""
    spans = _run_spans(para)
    if spans is None:
        return False
    old_text = "".join(s[2] for s in spans)
    if not old_text.strip():
        return False

    segments = _diff_segments(old_text, new_text)
    if not any(kind == "change" for kind, *_ in segments):
        return True  # new text is the old text; nothing to redline

    p = para._p
    index = list(p).index(spans[0][0])
    for el, *_ in spans:
        p.remove(el)

    nodes = []
    for kind, start, end, inserted in segments:
        pieces = _slice_spans(spans, start, end)
        if kind == "equal":
            nodes.extend(_make_text_run(text, rpr) for rpr, text in pieces)
            continue
        if pieces:
            deletion = _revision("w:del", author, date)
            for rpr, text in pieces:
                deletion.append(_make_del_run(text, rpr))
            nodes.append(deletion)
        if inserted:
            rpr = pieces[0][0] if pieces else _rpr_at(spans, start)
            insertion = _revision("w:ins", author, date)
            insertion.append(_make_text_run(inserted, rpr))
            nodes.append(insertion)

    for offset, node in enumerate(nodes):
        p.insert(index + offset, node)
    return True


def _rpr_at(spans, pos: int):
    """Run properties in force at character *pos* — for text inserted at a boundary."""
    for _el, rpr, _text, s, e in spans:
        if s < pos <= e:
            return rpr
    return spans[0][1]


def _replace_whole_tracked(para: Paragraph, new_text: str, author: str, date: str) -> None:
    """Redline a paragraph wholesale: strike every run, insert *new_text* after them."""
    rpr = _first_run_rpr(para)
    for run in list(para.runs):
        _mark_run_deleted(run, _revision("w:del", author, date))
    ins = _revision("w:ins", author, date)
    ins.append(_make_text_run(new_text, rpr))
    para._p.append(ins)


def _replace_tracked(
    para: Paragraph, new_text: str, author: str, date: str, word_level: bool = True
) -> str:
    """Redline a replacement; return the granularity used (``word`` or ``paragraph``)."""
    if word_level and _replace_word_tracked(para, new_text, author, date):
        return "word"
    _replace_whole_tracked(para, new_text, author, date)
    return "paragraph"


def _delete_tracked(para: Paragraph, author: str, date: str) -> None:
    """Redline a whole-paragraph deletion: strike every run and the paragraph mark."""
    for run in list(para.runs):
        _mark_run_deleted(run, _revision("w:del", author, date))
    _mark_para_mark(para, "w:del", author, date)


def _fill_inserted_paragraph(para: Paragraph, new_text: str, author: str, date: str) -> None:
    """Populate an empty new paragraph as a tracked insertion (runs + paragraph mark)."""
    ins = _revision("w:ins", author, date)
    ins.append(_make_text_run(new_text))
    para._p.append(ins)
    _mark_para_mark(para, "w:ins", author, date)


def _insert_after_tracked(para: Paragraph, new_text: str, author: str, date: str) -> Paragraph:
    """Insert a new tracked-insertion paragraph immediately after *para*."""
    new_p = OxmlElement("w:p")
    para._p.addnext(new_p)
    new_para = Paragraph(new_p, para._parent)
    try:
        new_para.style = para.style
    except Exception:
        pass  # style may not be assignable on some documents; leave default
    _fill_inserted_paragraph(new_para, new_text, author, date)
    return new_para


def _replace_clean(para: Paragraph, new_text: str) -> None:
    """Overwrite a paragraph's text directly, reusing the first run's formatting."""
    rpr = _first_run_rpr(para)
    for run in list(para.runs):
        run._element.getparent().remove(run._element)
    para._p.append(_make_text_run(new_text, rpr))


def _delete_clean(para: Paragraph) -> None:
    """Remove a whole paragraph outright."""
    para._p.getparent().remove(para._p)


def _insert_after_clean(para: Paragraph, new_text: str) -> Paragraph:
    """Insert a new plain paragraph immediately after *para*."""
    new_p = OxmlElement("w:p")
    para._p.addnext(new_p)
    new_para = Paragraph(new_p, para._parent)
    try:
        new_para.style = para.style
    except Exception:
        pass  # style may not be assignable on some documents; leave default
    new_para._p.append(_make_text_run(new_text))
    return new_para


def apply_operations(
    doc,
    operations: list[dict],
    author: str,
    date: str,
    tracked: bool = True,
    word_level: bool = True,
) -> list[dict]:
    """Apply operations in order; return a per-op match report.

    With *tracked* (the default) each change is recorded as Word tracked changes;
    otherwise the edits are applied cleanly, producing a final document with no
    tracked-change markup. With *word_level* (the default) a tracked ``replace``
    strikes only the words that changed rather than the whole paragraph.
    """
    report: list[dict] = []
    for op in operations:
        kind = op.get("op")
        anchor = op.get("anchor_text", "")
        new_text = op.get("new_text", "")
        matched = True
        granularity = None

        if kind == "append":
            if tracked:
                _fill_inserted_paragraph(doc.add_paragraph(), new_text, author, date)
            else:
                doc.add_paragraph(new_text)
        else:
            para = _find_paragraph(doc, anchor)
            if para is None:
                matched = False
            elif kind == "replace":
                if tracked:
                    granularity = _replace_tracked(para, new_text, author, date, word_level)
                else:
                    _replace_clean(para, new_text)
            elif kind == "insert_after":
                _insert_after_tracked(para, new_text, author, date) if tracked else _insert_after_clean(para, new_text)
            elif kind == "delete":
                _delete_tracked(para, author, date) if tracked else _delete_clean(para)
            else:
                matched = False  # unknown op type

        entry = {"op": kind, "anchor_text": anchor, "matched": matched}
        if granularity:
            entry["granularity"] = granularity
        report.append(entry)
    return report


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--original", required=True, help="path to the original .docx")
    parser.add_argument("--ops", required=True, help="path to the operations JSON file")
    parser.add_argument("--output", required=True, help="path to write the edited .docx")
    parser.add_argument(
        "--author",
        default="edit-docx",
        help="author name recorded on each tracked change (default: edit-docx)",
    )
    parser.add_argument(
        "--clean",
        action="store_true",
        help="apply edits directly (final document) instead of as tracked changes",
    )
    parser.add_argument(
        "--whole-paragraph",
        action="store_true",
        help="redline a replace by striking the whole paragraph instead of word-diffing it",
    )
    args = parser.parse_args()

    with open(args.ops, encoding="utf-8") as f:
        operations = json.load(f).get("operations", [])

    date = datetime.datetime.now(datetime.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    doc = Document(args.original)
    report = apply_operations(
        doc,
        operations,
        args.author,
        date,
        tracked=not args.clean,
        word_level=not args.whole_paragraph,
    )
    doc.save(args.output)

    unmatched = sum(1 for r in report if not r["matched"])
    json.dump({"operations": report, "unmatched": unmatched}, sys.stdout)
    return 0


if __name__ == "__main__":
    sys.exit(main())
