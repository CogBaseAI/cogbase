#!/usr/bin/env python3
"""Apply a list of amendment operations to an original .docx, preserving formatting.

The deterministic *apply* helper for the merge-contract skill. The agent does the
*understand* step (reading the amendment and deriving the operation list); this
script applies them, editing the original OOXML in place — matching each operation
to a paragraph by an anchor snippet and doing run-level edits — so styles,
numbering, and fonts survive.

Usage::

    python apply_operations.py --original in.docx --ops ops.json --output out.docx

``ops.json`` shape::

    {"operations": [
        {"op": "replace",      "anchor_text": "...", "new_text": "..."},
        {"op": "delete",       "anchor_text": "..."},
        {"op": "insert_after", "anchor_text": "...", "new_text": "..."},
        {"op": "append",                             "new_text": "..."}
    ]}

Prints a JSON report to stdout: per-operation ``matched`` flag plus an
``unmatched`` count, so the caller can surface operations that found no anchor
(a common sign the amendment referenced a section the original doesn't contain).
"""

from __future__ import annotations

import argparse
import json
import re
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def _norm(text: str) -> str:
    """Collapse whitespace and lowercase so anchor matching tolerates reflow."""
    return re.sub(r"\s+", " ", text or "").strip().lower()


def _find_paragraph(doc, anchor_text: str):
    """Return the first paragraph whose normalized text contains the anchor, or None."""
    needle = _norm(anchor_text)
    if not needle:
        return None
    for para in doc.paragraphs:
        if needle in _norm(para.text):
            return para
    return None


def _set_text_preserving_format(para: Paragraph, new_text: str) -> None:
    """Replace a paragraph's text while keeping its first run's formatting.

    Word stores formatting on runs; we keep the first run (its font/style) and
    retarget its text, then drop the remaining runs so no stale text lingers.
    """
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        para.add_run(new_text)


def _insert_paragraph_after(para: Paragraph, new_text: str) -> Paragraph:
    """Insert a new paragraph immediately after *para*, inheriting its style."""
    new_p = OxmlElement("w:p")
    para._p.addnext(new_p)
    new_para = Paragraph(new_p, para._parent)
    new_para.add_run(new_text)
    try:
        new_para.style = para.style
    except Exception:
        pass  # style may not be assignable on some documents; leave default
    return new_para


def _delete_paragraph(para: Paragraph) -> None:
    para._element.getparent().remove(para._element)


def apply_operations(doc, operations: list[dict]) -> list[dict]:
    """Apply operations in order; return a per-operation report with match flags."""
    report: list[dict] = []
    for op in operations:
        kind = op.get("op")
        anchor = op.get("anchor_text", "")
        new_text = op.get("new_text", "")
        matched = True

        if kind == "append":
            doc.add_paragraph(new_text)
        else:
            para = _find_paragraph(doc, anchor)
            if para is None:
                matched = False
            elif kind == "replace":
                _set_text_preserving_format(para, new_text)
            elif kind == "insert_after":
                _insert_paragraph_after(para, new_text)
            elif kind == "delete":
                _delete_paragraph(para)
            else:
                matched = False  # unknown op type

        report.append({"op": kind, "anchor_text": anchor, "matched": matched})
    return report


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--original", required=True, help="path to the original .docx")
    parser.add_argument("--ops", required=True, help="path to the operations JSON file")
    parser.add_argument("--output", required=True, help="path to write the merged .docx")
    args = parser.parse_args()

    with open(args.ops, encoding="utf-8") as f:
        operations = json.load(f).get("operations", [])

    doc = Document(args.original)
    report = apply_operations(doc, operations)
    doc.save(args.output)

    unmatched = sum(1 for r in report if not r["matched"])
    json.dump({"operations": report, "unmatched": unmatched}, sys.stdout)
    return 0


if __name__ == "__main__":
    sys.exit(main())
