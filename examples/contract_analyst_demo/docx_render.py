"""Render a plain-text contract fixture as a Word .docx.

The contract fixtures live as plain text in ``saas_contracts.py``. This module
renders them to .docx on demand so the demo can ingest them as Word files
(parsed to markdown server-side) without committing any binaries to the repo.

Shared by:
  - ``demo.py`` — writes temp .docx files to upload via the REST client
  - ``gen_demos_json.py`` — embeds .docx bytes in the demo catalog so the UI
    Demos tab can ingest the fixtures as Word files
"""

from __future__ import annotations

import io
import pathlib

DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _build_document(text: str):
    """Build a python-docx ``Document`` from a plain-text contract.

    Blank-line-separated blocks become paragraphs; short all-caps lines
    (e.g. "PARTIES", "PURPOSE") are rendered as headings so the parsed
    markdown keeps a sensible structure.
    """
    from docx import Document  # optional dep; see the `demo` extra

    doc = Document()
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if "\n" not in block and block.isupper() and len(block) < 60:
            doc.add_heading(block.title(), level=1)
        else:
            doc.add_paragraph(block)
    return doc


def to_docx_bytes(text: str) -> bytes:
    """Render *text* to an in-memory .docx and return the raw bytes."""
    buf = io.BytesIO()
    _build_document(text).save(buf)
    return buf.getvalue()


def write_docx(text: str, path: pathlib.Path) -> None:
    """Render *text* to a .docx file at *path*."""
    _build_document(text).save(str(path))
