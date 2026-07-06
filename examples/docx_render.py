"""Render a plain-text demo fixture as a Word .docx.

The demo fixtures live as plain text in each demo's data module. This module
renders them to .docx on demand so the demos can ingest them as Word files
(parsed to markdown server-side) without committing any binaries to the repo.

Shared by:
  - each demo's ``demo.py`` — via ``CogBaseClient.upload_docx_documents``, which
    writes temp .docx files and uploads them through the REST client
  - ``gen_demos_json.py`` — embeds .docx bytes in the demo catalog so the UI
    Demos tab can ingest the fixtures as Word files
"""

from __future__ import annotations

import io
import pathlib

DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _build_document(text: str):
    """Build a python-docx ``Document`` from a plain-text contract.

    Blank-line-separated blocks become paragraphs; short all-caps lines
    (e.g. "PARTIES", "PURPOSE") are rendered as headings so the parsed
    markdown keeps a sensible structure.
    """
    from docx import Document  # optional dep; see the `demo` extra

    doc = Document()
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if "\n" not in block and block.isupper() and len(block) < 60:
            doc.add_heading(block.title(), level=1)
        else:
            doc.add_paragraph(block)
    return doc


def to_docx_bytes(text: str) -> bytes:
    """Render *text* to an in-memory .docx and return the raw bytes."""
    buf = io.BytesIO()
    _build_document(text).save(buf)
    return buf.getvalue()


def write_docx(text: str, path: pathlib.Path) -> None:
    """Render *text* to a .docx file at *path*."""
    _build_document(text).save(str(path))
