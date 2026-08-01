"""Unit tests for the edit-docx skill's apply_operations helper.

The helper is a standalone CLI shipped inside the skill bundle (not an importable
package), so it is loaded by path. Tests cover the paragraph-level apply logic
(replace / delete / insert_after / append), formatting preservation, the unmatched
report, and the CLI entry point end-to-end.

The helper produces a **redline**: edits are recorded as Word tracked changes
(``<w:ins>`` / ``<w:del>``), not clean overwrites. python-docx's ``paragraph.text``
and ``.runs`` do not see runs nested inside those wrappers, so the assertions here
inspect the tracked-change markup directly — inserted text from ``<w:ins>//<w:t>``
and deleted text from ``<w:del>//<w:delText>`` — rather than the visible text.

A tracked ``replace`` is word-diffed against the original paragraph, so only the
changed words are struck and re-inserted. The assertions that matter most are
``_accepted``/``_rejected``: whatever the diff chooses to mark up, accepting every
change in Word must yield exactly ``new_text`` and rejecting every change must
restore the original.

Skipped entirely when python-docx is not installed (it is a skill-declared
dependency, installed into the skill's venv at load time, not a test dependency).
"""

from __future__ import annotations

import importlib.util
import json
import pathlib
import subprocess
import sys

import pytest

pytest.importorskip("docx")

from docx import Document  # noqa: E402  (import after importorskip)
from docx.oxml.ns import qn  # noqa: E402

SKILL_DIR = pathlib.Path(__file__).resolve().parents[2] / "skills" / "edit-docx"
SCRIPT = SKILL_DIR / "apply_operations.py"

AUTHOR = "tester"
DATE = "2026-01-01T00:00:00Z"


def _load_helper():
    spec = importlib.util.spec_from_file_location("apply_operations", SCRIPT)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


helper = _load_helper()


def _apply(doc, operations):
    """apply_operations with fixed author/date so tests don't repeat them."""
    return helper.apply_operations(doc, operations, AUTHOR, DATE)


def _apply_clean(doc, operations):
    """apply_operations in clean mode — edits baked in, no tracked-change markup."""
    return helper.apply_operations(doc, operations, AUTHOR, DATE, tracked=False)


def _doc(paragraphs: list[str]) -> Document:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    return doc


# --- tracked-change readers -------------------------------------------------
# python-docx can't see runs inside <w:ins>/<w:del>, so read the markup directly.


def _inserted_text(scope) -> str:
    """Concatenated text of every ``<w:ins>//<w:t>`` under *scope* (doc or paragraph)."""
    return "".join(
        t.text or "" for ins in scope.iter(qn("w:ins")) for t in ins.iter(qn("w:t"))
    )


def _deleted_text(scope) -> str:
    """Concatenated text of every ``<w:del>//<w:delText>`` under *scope*."""
    return "".join(
        t.text or "" for d in scope.iter(qn("w:del")) for t in d.iter(qn("w:delText"))
    )


def _el(doc_or_para):
    """The lxml element to iterate: a paragraph's ``<w:p>`` or the document body."""
    return doc_or_para._p if hasattr(doc_or_para, "_p") else doc_or_para.element.body


def _ins_text(scope):
    return _inserted_text(_el(scope))


def _del_text(scope):
    return _deleted_text(_el(scope))


def _resolve(para, accept: bool) -> str:
    """Paragraph text as Word would render it after accept-all / reject-all.

    Walks the paragraph's direct children in document order — plain ``<w:r>`` is
    surviving text, ``<w:ins>`` counts only when accepting, ``<w:del>`` only when
    rejecting — which is exactly the round trip a reviewer performs.
    """
    out = []
    for child in para._p:
        tag = child.tag.split("}")[1]
        if tag == "r" or (tag == "ins" and accept) or (tag == "del" and not accept):
            out.append(child.xpath("string(.)"))
    return "".join(out)


def _accepted(para) -> str:
    return _resolve(para, accept=True)


def _rejected(para) -> str:
    return _resolve(para, accept=False)


def _para_mark_tag(para) -> str | None:
    """Return 'w:ins'/'w:del' if the paragraph mark itself is tracked, else None."""
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        return None
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        return None
    for tag in ("w:ins", "w:del"):
        if rPr.find(qn(tag)) is not None:
            return tag
    return None


# ---------------------------------------------------------------------------
# apply_operations — per-op behavior (as tracked changes)
# ---------------------------------------------------------------------------


def test_replace_redlines_only_the_changed_words():
    doc = _doc(["Payment shall be due within 30 days.", "Other clause."])
    report = _apply(
        doc,
        [{"op": "replace", "anchor_text": "Payment shall be due within 30 days",
          "new_text": "Payment shall be due within 45 days."}],
    )
    para = doc.paragraphs[0]
    # only "30" is struck and only "45" inserted — the rest of the sentence is untouched
    assert _del_text(para) == "30"
    assert _ins_text(para) == "45"
    # the unchanged wording stays live text, not markup
    assert "Payment shall be due within" in para.text
    # and the round trip is exact in both directions
    assert _accepted(para) == "Payment shall be due within 45 days."
    assert _rejected(para) == "Payment shall be due within 30 days."
    # the second (untouched) paragraph keeps its plain visible text
    assert doc.paragraphs[1].text == "Other clause."
    assert report == [{"op": "replace", "anchor_text": "Payment shall be due within 30 days",
                       "matched": True, "granularity": "word"}]


def test_replace_records_author_and_date():
    doc = _doc(["Payment shall be due within 30 days."])
    _apply(doc, [{"op": "replace", "anchor_text": "Payment shall be due within 30 days", "new_text": "x"}])
    ins = doc.paragraphs[0]._p.find(qn("w:ins"))
    assert ins.get(qn("w:author")) == AUTHOR
    assert ins.get(qn("w:date")) == DATE


def test_replace_preserves_run_formatting():
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("Payment due in 30 days")
    run.bold = True

    _apply(doc, [{"op": "replace", "anchor_text": "Payment due in 30 days", "new_text": "Payment due in 45 days"}])

    # the inserted run clones the original run's <w:rPr>, so bold survives the redline
    ins = para._p.find(qn("w:ins"))
    ins_run = ins.find(qn("w:r"))
    assert ins_run.find(qn("w:rPr")).find(qn("w:b")) is not None
    assert _ins_text(para) == "45"
    # the surviving halves of the split run keep the formatting too
    assert all(r.bold for r in para.runs)


def test_delete_strikes_paragraph_and_mark():
    doc = _doc(["Keep this.", "Either party may terminate with 60 days notice.", "Keep that."])
    report = _apply(
        doc, [{"op": "delete", "anchor_text": "Either party may terminate with 60 days notice"}]
    )
    # paragraph is not removed — it stays as a tracked deletion (text + paragraph mark)
    assert len(doc.paragraphs) == 3
    target = doc.paragraphs[1]
    assert "Either party may terminate with 60 days notice." in _del_text(target)
    assert _para_mark_tag(target) == "w:del"
    # surrounding paragraphs untouched
    assert doc.paragraphs[0].text == "Keep this."
    assert doc.paragraphs[2].text == "Keep that."
    assert report[0]["matched"] is True


def test_insert_after_adds_tracked_paragraph_in_position():
    doc = _doc(["Section 8 Term", "Section 10 Misc"])
    _apply(
        doc,
        [{"op": "insert_after", "anchor_text": "Section 8 Term",
          "new_text": "Section 9 Governing Law: State of Delaware."}],
    )
    assert len(doc.paragraphs) == 3
    inserted = doc.paragraphs[1]
    assert _ins_text(inserted) == "Section 9 Governing Law: State of Delaware."
    assert _para_mark_tag(inserted) == "w:ins"
    # placed between the anchor and the following paragraph
    assert doc.paragraphs[0].text == "Section 8 Term"
    assert doc.paragraphs[2].text == "Section 10 Misc"


def test_append_adds_tracked_paragraph_at_end():
    doc = _doc(["First.", "Second."])
    _apply(doc, [{"op": "append", "new_text": "Appended clause."}])
    last = doc.paragraphs[-1]
    assert _ins_text(last) == "Appended clause."
    assert _para_mark_tag(last) == "w:ins"


def test_anchor_matching_is_whitespace_and_case_insensitive():
    doc = _doc(["Section   4.2   PAYMENT terms apply."])
    report = _apply(
        doc, [{"op": "replace", "anchor_text": "section 4.2 payment terms", "new_text": "Replaced."}]
    )
    assert report[0]["matched"] is True
    assert _accepted(doc.paragraphs[0]) == "Replaced."


# ---------------------------------------------------------------------------
# markdown-tolerant anchor matching
#
# The agent reads the base as markdown (docx is extracted to markdown at ingest),
# so it copies anchors containing `**bold**`, `3.` list prefixes, and `#` headings.
# The docx paragraph text is raw — no markdown. Matching must bridge that gap so a
# verbatim-from-markdown anchor still locates the raw paragraph on the first pass.
# ---------------------------------------------------------------------------


def test_norm_strips_inline_emphasis_and_leading_markers():
    # leading list number + inline bold in the anchor; neither in the raw paragraph
    assert helper._norm("3. **Security Deposit:** pay **$4500** now") == \
        helper._norm("Security Deposit: pay $4500 now")


def test_replace_matches_markdown_anchor_against_raw_paragraph():
    # Reproduces the log.1 failure: the agent's anchor carried markdown, the docx
    # paragraph did not, so the anchor never matched and the agent burned its budget.
    doc = _doc(["Security Deposit: Tenant shall pay a security deposit of $4500 to Landlord."])
    report = _apply(
        doc,
        [{"op": "replace",
          "anchor_text": "3. **Security Deposit:** Tenant shall pay a security deposit of **$4500** to Landlord.",
          "new_text": "Security Deposit: Tenant shall pay a security deposit of $5000 to Landlord."}],
    )
    assert report[0]["matched"] is True
    # the anchor's markdown is stripped for matching; the diff narrows to the amount
    assert _del_text(doc) == "4500"
    assert _ins_text(doc) == "5000"


def test_delete_matches_markdown_heading_anchor():
    doc = _doc(["Confidentiality", "Keep this."])
    report = _apply(
        doc, [{"op": "delete", "anchor_text": "## Confidentiality"}]
    )
    assert report[0]["matched"] is True
    assert "Confidentiality" in _del_text(doc)
    assert _para_mark_tag(doc.paragraphs[0]) == "w:del"
    assert doc.paragraphs[1].text == "Keep this."


def test_insert_after_matches_bulleted_markdown_anchor():
    doc = _doc(["Provider shall notify Customer within 48 hours.", "Section 6."])
    _apply(
        doc,
        [{"op": "insert_after",
          "anchor_text": "- Provider shall notify Customer within **48 hours**.",
          "new_text": "5.5 Subprocessors clause."}],
    )
    assert len(doc.paragraphs) == 3
    assert _ins_text(doc.paragraphs[1]) == "5.5 Subprocessors clause."
    assert doc.paragraphs[0].text == "Provider shall notify Customer within 48 hours."
    assert doc.paragraphs[2].text == "Section 6."


# ---------------------------------------------------------------------------
# word-level replace redlines
#
# A tracked replace is diffed word-by-word against the original paragraph so the
# reviewer sees "within ~~30~~45 days" instead of a struck sentence followed by a
# near-identical one. The invariant under all of it: accept-all == new_text,
# reject-all == the original.
# ---------------------------------------------------------------------------


WORD_DIFF_CASES = [
    # (original, replacement)
    ("Payment shall be due within 30 days of invoice receipt.",
     "Payment shall be due within 45 days of invoice receipt."),
    ("Either party may terminate with 60 days notice.",
     "Either party may terminate with 90 days written notice to the other party."),
    ("The Client shall indemnify the Provider for all claims whatsoever.",
     "The Client shall indemnify the Provider for direct claims."),          # deletion-heavy
    ("Confidentiality survives termination.",
     "Confidentiality survives termination for three (3) years."),           # pure addition
    ("Governing law: New York.",
     "This Agreement is governed by the laws of Delaware."),                 # full rewrite
    ("Fees are $1,000.00 per month.", "Fees are $1,200.50 per month."),
    ("乙方应当在合同终止后三年内承担保密义务。", "乙方应当在合同终止后两年内承担保密义务。"),
    ("A B C D E F G", "A X C D Y F G"),                                      # scattered edits
    ("Unchanged clause.", "Unchanged clause."),                              # no-op
]


@pytest.mark.parametrize("original,replacement", WORD_DIFF_CASES)
@pytest.mark.parametrize("word_level", [True, False])
def test_replace_round_trips_through_accept_and_reject(original, replacement, word_level):
    doc = _doc([original])
    helper.apply_operations(
        doc,
        [{"op": "replace", "anchor_text": original, "new_text": replacement}],
        AUTHOR, DATE, word_level=word_level,
    )
    para = doc.paragraphs[0]
    assert _accepted(para) == replacement
    assert _rejected(para) == original


def test_word_diff_marks_up_less_than_whole_paragraph_replace():
    original = "Either party may terminate this Agreement with 60 days notice."
    replacement = "Either party may terminate this Agreement with 90 days notice."
    ops = [{"op": "replace", "anchor_text": original, "new_text": replacement}]

    narrow, wide = _doc([original]), _doc([original])
    helper.apply_operations(narrow, ops, AUTHOR, DATE, word_level=True)
    helper.apply_operations(wide, ops, AUTHOR, DATE, word_level=False)

    assert _del_text(narrow) == "60"
    assert _del_text(wide) == original  # the pre-word-diff behavior, still available


def test_word_diff_absorbs_a_stranded_short_equal_run():
    # "30 days" -> "45 business days" leaves " " equal between two changes; merging it
    # keeps the markup one edit instead of two strikeouts around a stranded space
    doc = _doc(["Notice is due within 30 days."])
    _apply(doc, [{"op": "replace", "anchor_text": "Notice is due",
                  "new_text": "Notice is due within 45 business days."}])
    para = doc.paragraphs[0]
    assert len(para._p.findall(qn("w:del"))) == 1
    assert len(para._p.findall(qn("w:ins"))) == 1
    assert _accepted(para) == "Notice is due within 45 business days."


def test_word_diff_of_unspaced_script_narrows_to_characters():
    doc = _doc(["乙方应当在合同终止后三年内对甲方的商业秘密承担保密义务。"])
    _apply(doc, [{"op": "replace", "anchor_text": "承担保密义务",
                  "new_text": "乙方应当在合同终止后两年内对甲方的商业秘密承担保密义务。"}])
    # CJK has no spaces to split on, so tokens are single characters
    assert _del_text(doc) == "三"
    assert _ins_text(doc) == "两"


def test_pure_addition_produces_no_deletion():
    doc = _doc(["Confidentiality survives termination."])
    _apply(doc, [{"op": "replace", "anchor_text": "Confidentiality survives",
                  "new_text": "Confidentiality survives termination for three (3) years."}])
    assert _del_text(doc) == ""
    assert _ins_text(doc) == " for three (3) years"


def test_identical_replacement_produces_no_markup():
    doc = _doc(["Unchanged clause."])
    report = _apply(doc, [{"op": "replace", "anchor_text": "Unchanged clause",
                           "new_text": "Unchanged clause."}])
    assert report[0]["granularity"] == "word"
    assert _ins_text(doc) == "" and _del_text(doc) == ""
    assert doc.paragraphs[0].text == "Unchanged clause."


def test_word_diff_preserves_formatting_across_an_edit_boundary():
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("5.2 Liability. ")
    bold = para.add_run("liability shall not exceed twelve (12) months of fees")
    bold.bold = True
    para.add_run(", except for gross negligence.")

    _apply(doc, [{"op": "replace", "anchor_text": "liability shall not exceed",
                  "new_text": "5.2 Liability. liability shall not exceed twenty-four (24) months"
                              " of fees, except for gross negligence."}])

    # the bold run is split around the edit; both halves — and the insertion — stay bold
    assert _ins_text(para).strip().startswith("twenty-four")
    ins_run = para._p.find(qn("w:ins")).find(qn("w:r"))
    assert ins_run.find(qn("w:rPr")).find(qn("w:b")) is not None
    assert "5.2 Liability. " in para.text and ", except for gross negligence." in para.text


def test_paragraph_with_a_tab_falls_back_to_whole_paragraph():
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("8. Term.\t")
    para.add_run("This Agreement runs for one year.")

    report = _apply(doc, [{"op": "replace", "anchor_text": "This Agreement runs for one year",
                           "new_text": "8. Term.\tThis Agreement runs for two years."}])

    # a tab isn't sliceable by character offset, so the paragraph is struck whole
    assert report[0]["granularity"] == "paragraph"
    assert "This Agreement runs for one year." in _del_text(para)
    assert _ins_text(para) == "8. Term.\tThis Agreement runs for two years."


def test_cli_whole_paragraph_flag_disables_the_word_diff(tmp_path):
    original = tmp_path / "in.docx"
    _doc(["Payment shall be due within 30 days."]).save(str(original))
    ops_path = tmp_path / "ops.json"
    ops_path.write_text(json.dumps({"operations": [
        {"op": "replace", "anchor_text": "Payment shall be due within 30 days",
         "new_text": "Payment shall be due within 45 days."}]}))
    output = tmp_path / "out.docx"

    proc = subprocess.run(
        [sys.executable, str(SCRIPT), "--original", str(original), "--ops", str(ops_path),
         "--output", str(output), "--whole-paragraph"],
        capture_output=True, text=True,
    )
    assert proc.returncode == 0, proc.stderr
    assert json.loads(proc.stdout)["operations"][0]["granularity"] == "paragraph"
    assert _deleted_text(Document(str(output)).element.body) == "Payment shall be due within 30 days."


# ---------------------------------------------------------------------------
# unmatched / unknown operations
# ---------------------------------------------------------------------------


def test_unmatched_anchor_is_reported_not_applied():
    doc = _doc(["Only clause."])
    report = _apply(
        doc, [{"op": "replace", "anchor_text": "nonexistent section", "new_text": "x"}]
    )
    assert report[0]["matched"] is False
    assert doc.paragraphs[0].text == "Only clause."  # unchanged
    assert _ins_text(doc) == "" and _del_text(doc) == ""  # no redline produced


def test_unknown_op_reported_unmatched():
    doc = _doc(["Clause."])
    report = _apply(doc, [{"op": "frobnicate", "anchor_text": "Clause"}])
    assert report[0]["matched"] is False


# ---------------------------------------------------------------------------
# CLI end-to-end
# ---------------------------------------------------------------------------


def test_cli_applies_ops_and_reports(tmp_path):
    original = tmp_path / "in.docx"
    _doc(["Payment shall be due within 30 days.", "Termination clause here."]).save(str(original))

    ops = {"operations": [
        {"op": "replace", "anchor_text": "Payment shall be due within 30 days",
         "new_text": "Payment shall be due within 45 days."},
        {"op": "delete", "anchor_text": "Termination clause here"},
        {"op": "append", "new_text": "Governing law: Delaware."},
        {"op": "insert_after", "anchor_text": "does not exist", "new_text": "y"},
    ]}
    ops_path = tmp_path / "ops.json"
    ops_path.write_text(json.dumps(ops))
    output = tmp_path / "out.docx"

    proc = subprocess.run(
        [sys.executable, str(SCRIPT),
         "--original", str(original), "--ops", str(ops_path), "--output", str(output),
         "--author", "Jun Luo"],
        capture_output=True, text=True,
    )
    assert proc.returncode == 0, proc.stderr

    report = json.loads(proc.stdout)
    assert report["unmatched"] == 1
    assert [r["matched"] for r in report["operations"]] == [True, True, True, False]
    assert report["operations"][0]["granularity"] == "word"

    redlined = Document(str(output))
    inserted = _inserted_text(redlined.element.body)
    deleted = _deleted_text(redlined.element.body)
    # the replace narrows to the changed term; append/delete stay whole-paragraph
    assert _accepted(redlined.paragraphs[0]) == "Payment shall be due within 45 days."
    assert _rejected(redlined.paragraphs[0]) == "Payment shall be due within 30 days."
    assert "45" in inserted and "30" in deleted
    assert "Governing law: Delaware." in inserted
    assert "Termination clause here." in deleted
    # author propagated from the CLI flag onto the tracked changes
    assert redlined.element.body.iter(qn("w:ins")).__next__().get(qn("w:author")) == "Jun Luo"


# ---------------------------------------------------------------------------
# clean mode (tracked=False / --clean) — edits baked in as a final document
#
# The legal-review skill hands back a clean final docx once the user has
# accepted/rejected the redline's suggestions: accepted changes applied directly,
# with no <w:ins>/<w:del> markup. Here the edits land in the *visible* text, so
# these assertions read paragraph.text/.runs and confirm no tracked changes exist.
# ---------------------------------------------------------------------------


def _has_tracked_changes(doc) -> bool:
    body = doc.element.body
    return (
        next(body.iter(qn("w:ins")), None) is not None
        or next(body.iter(qn("w:del")), None) is not None
    )


def test_clean_replace_overwrites_visible_text():
    doc = _doc(["Payment shall be due within 30 days.", "Other clause."])
    report = _apply_clean(
        doc,
        [{"op": "replace", "anchor_text": "Payment shall be due within 30 days",
          "new_text": "Payment shall be due within 45 days."}],
    )
    # new text is live; old text is gone entirely — not struck through
    assert doc.paragraphs[0].text == "Payment shall be due within 45 days."
    assert doc.paragraphs[1].text == "Other clause."
    assert not _has_tracked_changes(doc)
    assert report == [{"op": "replace", "anchor_text": "Payment shall be due within 30 days", "matched": True}]


def test_clean_replace_preserves_run_formatting():
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("Payment due in 30 days")
    run.bold = True

    _apply_clean(doc, [{"op": "replace", "anchor_text": "Payment due in 30 days", "new_text": "Payment due in 45 days"}])

    assert para.text == "Payment due in 45 days"
    # the surviving run clones the original <w:rPr>, so bold carries over
    assert para.runs[0].bold is True
    assert not _has_tracked_changes(doc)


def test_clean_delete_removes_paragraph():
    doc = _doc(["Keep this.", "Either party may terminate with 60 days notice.", "Keep that."])
    report = _apply_clean(
        doc, [{"op": "delete", "anchor_text": "Either party may terminate with 60 days notice"}]
    )
    # paragraph is gone outright — not retained as a struck-through deletion
    assert [p.text for p in doc.paragraphs] == ["Keep this.", "Keep that."]
    assert not _has_tracked_changes(doc)
    assert report[0]["matched"] is True


def test_clean_insert_after_adds_plain_paragraph_in_position():
    doc = _doc(["Section 8 Term", "Section 10 Misc"])
    _apply_clean(
        doc,
        [{"op": "insert_after", "anchor_text": "Section 8 Term",
          "new_text": "Section 9 Governing Law: State of Delaware."}],
    )
    assert [p.text for p in doc.paragraphs] == [
        "Section 8 Term",
        "Section 9 Governing Law: State of Delaware.",
        "Section 10 Misc",
    ]
    assert not _has_tracked_changes(doc)


def test_clean_append_adds_plain_paragraph_at_end():
    doc = _doc(["First.", "Second."])
    _apply_clean(doc, [{"op": "append", "new_text": "Appended clause."}])
    assert doc.paragraphs[-1].text == "Appended clause."
    assert not _has_tracked_changes(doc)


def test_clean_unmatched_anchor_is_reported_not_applied():
    doc = _doc(["Only clause."])
    report = _apply_clean(
        doc, [{"op": "replace", "anchor_text": "nonexistent section", "new_text": "x"}]
    )
    assert report[0]["matched"] is False
    assert doc.paragraphs[0].text == "Only clause."  # unchanged


def test_cli_clean_flag_bakes_in_accepted_changes(tmp_path):
    original = tmp_path / "in.docx"
    _doc(["Payment shall be due within 30 days.", "Termination clause here."]).save(str(original))

    # mirrors the legal-review "final docx" pass: only accepted ops, applied clean
    ops = {"operations": [
        {"op": "replace", "anchor_text": "Payment shall be due within 30 days",
         "new_text": "Payment shall be due within 45 days."},
        {"op": "delete", "anchor_text": "Termination clause here"},
    ]}
    ops_path = tmp_path / "ops.json"
    ops_path.write_text(json.dumps(ops))
    output = tmp_path / "final.docx"

    proc = subprocess.run(
        [sys.executable, str(SCRIPT),
         "--original", str(original), "--ops", str(ops_path), "--output", str(output),
         "--clean"],
        capture_output=True, text=True,
    )
    assert proc.returncode == 0, proc.stderr

    report = json.loads(proc.stdout)
    assert report["unmatched"] == 0
    assert [r["matched"] for r in report["operations"]] == [True, True]

    final = Document(str(output))
    # accepted change is live; deleted clause is gone; no redline markup remains
    assert [p.text for p in final.paragraphs] == ["Payment shall be due within 45 days."]
    assert not _has_tracked_changes(final)
