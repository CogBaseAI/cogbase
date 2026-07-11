"""Unit tests for the contract-review skill's deterministic helpers.

Two standalone CLIs shipped inside the skill bundle (loaded by path, like the
edit-docx tests):

  segment_clauses.py — .docx -> clauses with per-paragraph verbatim anchors
  build_ops.py       — finalize (raw analysis -> review ops.json) and
                       to-edit-ops (review -> edit-docx operations)

The suite also asserts the *contract* between the two skills: an anchor baked by
``finalize`` (taken verbatim from segmentation) is located by edit-docx's own
paragraph matcher — the property that makes the redline reliable.

Skipped when python-docx is not installed (a skill-declared dependency).
"""

from __future__ import annotations

import importlib.util
import pathlib

import pytest

pytest.importorskip("docx")

from docx import Document  # noqa: E402

SKILLS = pathlib.Path(__file__).resolve().parents[2] / "skills"
CR_DIR = SKILLS / "contract-review"


def _load(name: str, path: pathlib.Path):
    spec = importlib.util.spec_from_file_location(name, path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


segment_mod = _load("segment_clauses", CR_DIR / "segment_clauses.py")
build_mod = _load("build_ops", CR_DIR / "build_ops.py")
edit_helper = _load("apply_operations", SKILLS / "edit-docx" / "apply_operations.py")


def _contract_doc() -> Document:
    doc = Document()
    doc.add_heading("MASTER SERVICES AGREEMENT", level=0)
    doc.add_paragraph("1. PAYMENT TERMS")
    doc.add_paragraph("Payment shall be due within 30 days of invoice receipt.")
    doc.add_paragraph("2. TERMINATION")
    doc.add_paragraph("Either party may terminate with 60 days written notice.")
    return doc


# ---------------------------------------------------------------------------
# segment_clauses
# ---------------------------------------------------------------------------


def test_segment_groups_paragraphs_into_clauses_at_headings():
    clauses = segment_mod.segment(_contract_doc())
    headings = [c["heading"] for c in clauses]
    assert "1. PAYMENT TERMS" in headings
    assert "2. TERMINATION" in headings
    # The payment clause carries both its heading and its body paragraph.
    payment = next(c for c in clauses if c["heading"] == "1. PAYMENT TERMS")
    texts = [p["text"] for p in payment["paragraphs"]]
    assert "Payment shall be due within 30 days of invoice receipt." in texts


def test_segment_para_ids_are_stable_and_unique():
    clauses = segment_mod.segment(_contract_doc())
    ids = [p["para_id"] for c in clauses for p in c["paragraphs"]]
    assert len(ids) == len(set(ids))
    for c in clauses:
        for i, p in enumerate(c["paragraphs"]):
            assert p["para_id"] == f"{c['clause_id']}.p{i}"


def test_segment_skips_empty_paragraphs():
    doc = Document()
    doc.add_paragraph("1. TERM")
    doc.add_paragraph("")
    doc.add_paragraph("The term is one year.")
    clauses = segment_mod.segment(doc)
    texts = [p["text"] for c in clauses for p in c["paragraphs"]]
    assert "" not in texts


def test_segment_preamble_before_first_heading_is_clause_c0():
    doc = Document()
    doc.add_paragraph("This preamble has no section number.")
    doc.add_paragraph("1. SCOPE")
    doc.add_paragraph("Scope text.")
    clauses = segment_mod.segment(doc)
    assert clauses[0]["clause_id"] == "c0"
    assert clauses[0]["paragraphs"][0]["text"].startswith("This preamble")


# ---------------------------------------------------------------------------
# build_ops.finalize
# ---------------------------------------------------------------------------


def _clauses_fixture() -> list[dict]:
    return [
        {"clause_id": "c1", "heading": "1. PAYMENT TERMS", "paragraphs": [
            {"para_id": "c1.p0", "text": "1. PAYMENT TERMS"},
            {"para_id": "c1.p1", "text": "Payment shall be due within 30 days of invoice receipt."},
        ]},
    ]


def test_finalize_resolves_para_id_to_verbatim_anchor():
    analysis = {
        "base_doc_id": "d1",
        "meta": {"review_position": "disadvantaged"},
        "analyses": [
            {"clause_id": "c1", "risk": {"level": "high", "rationale": "tight"},
             "suggestion": {"op": "replace", "para_id": "c1.p1",
                            "new_text": "Payment shall be due within 45 days of invoice receipt."}},
        ],
    }
    review = build_mod.finalize(analysis, _clauses_fixture())
    clause = review["clauses"][0]
    assert clause["suggestion"]["anchor_text"] == "Payment shall be due within 30 days of invoice receipt."
    assert clause["verdict"] == "pending"
    assert clause["heading"] == "1. PAYMENT TERMS"


def test_finalize_keeps_clauses_without_a_suggestion():
    analysis = {"analyses": [{"clause_id": "c1", "risk": {"level": "none", "rationale": ""},
                              "suggestion": None}]}
    review = build_mod.finalize(analysis, _clauses_fixture())
    assert review["clauses"][0]["suggestion"] is None


def test_finalize_rejects_unknown_para_id():
    analysis = {"analyses": [{"clause_id": "c1",
                              "suggestion": {"op": "replace", "para_id": "c9.p9", "new_text": "x"}}]}
    with pytest.raises(ValueError, match="not found in segmentation"):
        build_mod.finalize(analysis, _clauses_fixture())


def test_finalize_rejects_bad_review_position():
    analysis = {"meta": {"review_position": "aggressive"}, "analyses": []}
    with pytest.raises(ValueError, match="review_position"):
        build_mod.finalize(analysis, _clauses_fixture())


def test_finalize_rejects_op_missing_required_text():
    analysis = {"analyses": [{"clause_id": "c1",
                              "suggestion": {"op": "replace", "para_id": "c1.p1"}}]}
    with pytest.raises(ValueError, match="requires new_text"):
        build_mod.finalize(analysis, _clauses_fixture())


def test_finalize_append_needs_no_anchor():
    analysis = {"analyses": [{"clause_id": "c1",
                              "suggestion": {"op": "append", "new_text": "Signed."}}]}
    review = build_mod.finalize(analysis, _clauses_fixture())
    sug = review["clauses"][0]["suggestion"]
    assert sug == {"op": "append", "new_text": "Signed."}


# ---------------------------------------------------------------------------
# build_ops.to_edit_ops
# ---------------------------------------------------------------------------


def _review_fixture() -> dict:
    return {"clauses": [
        {"clause_id": "c1", "verdict": "accepted",
         "suggestion": {"op": "replace", "anchor_text": "A", "new_text": "A2"}},
        {"clause_id": "c2", "verdict": "rejected",
         "suggestion": {"op": "replace", "anchor_text": "B", "new_text": "B2"}},
        {"clause_id": "c3", "verdict": "pending",
         "suggestion": {"op": "delete", "anchor_text": "C"}},
        {"clause_id": "c4", "verdict": "accepted", "suggestion": None},
    ]}


def test_to_edit_ops_accepted_only_by_default():
    ops = build_mod.to_edit_ops(_review_fixture())["operations"]
    assert ops == [{"op": "replace", "anchor_text": "A", "new_text": "A2"}]


def test_to_edit_ops_all_includes_every_suggestion_for_preview():
    ops = build_mod.to_edit_ops(_review_fixture(), accepted_only=False)["operations"]
    anchors = [o.get("anchor_text") for o in ops]
    assert anchors == ["A", "B", "C"]  # the null-suggestion clause is skipped


# ---------------------------------------------------------------------------
# cross-skill contract: baked anchors match edit-docx's paragraph finder
# ---------------------------------------------------------------------------


def test_baked_anchor_locates_paragraph_in_edit_docx():
    clauses = segment_mod.segment(_contract_doc())
    # Locate the termination body paragraph and target it by para_id.
    term_pid = next(
        p["para_id"] for c in clauses for p in c["paragraphs"]
        if "60 days" in p["text"]
    )
    analysis = {"analyses": [{"clause_id": "c2",
                              "suggestion": {"op": "replace", "para_id": term_pid,
                                             "new_text": "Either party may terminate with 90 days written notice."}}]}
    review = build_mod.finalize(analysis, clauses)
    anchor = review["clauses"][0]["suggestion"]["anchor_text"]

    # edit-docx's own matcher must find that anchor in the same base document.
    doc = _contract_doc()
    assert edit_helper._find_paragraph(doc, anchor) is not None
