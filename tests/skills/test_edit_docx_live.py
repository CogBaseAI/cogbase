"""Live end-to-end test for the edit-docx skill.

Nothing below the LLM boundary is faked. A real ``LocalFSDocumentStore`` backs the
document tools, the skill is loaded from disk with ``load_skill_dir`` — which
installs its declared ``python-docx`` dependency into the skill's private venv —
and the ``QueryRunner`` drives the real agent loop: routing to the skill, reading
the amendment, fetching the original ``.docx``, shelling out to the bundled
``apply_operations.py``, and saving the merged artifact.

The test builds a small contract and a matching amendment as real ``.docx`` files,
seeds them the way an upload would (extracted text + original bytes), runs the
merge through the runner, then re-opens the produced ``.docx`` and asserts the
amendment's three edits landed as a **redline** (Word tracked changes) — a ``replace``
(payment term: old struck, new inserted), a ``delete`` (the termination clause struck),
and an ``insert`` (a new governing-law section inserted) — while an untouched clause
survives as plain, non-tracked text.

Run with::

    pytest tests/skills/test_edit_docx_live.py -m live
"""

from __future__ import annotations

import contextlib
import glob
import os
from io import BytesIO
from pathlib import Path

import pytest

pytest.importorskip("docx")  # needed to build fixtures and read the merged output

from docx import Document  # noqa: E402  (import after importorskip)
from docx.oxml.ns import qn  # noqa: E402

from cogbase.core.query_runner import (  # noqa: E402
    MemoryTiers,
    QueryRunner,
    RetrievalResources,
)
from cogbase.skills.skill import load_skill_dir  # noqa: E402
from cogbase.stores.document.local_fs import LocalFSDocumentStore  # noqa: E402
from tests.live_setup import make_llm  # noqa: E402

_llm = make_llm()

pytestmark = [
    pytest.mark.live,
    pytest.mark.skipif(_llm is None, reason="No LLM configured (see tests/live_setup)"),
]

_SKILL_DIR = Path(__file__).resolve().parents[2] / "skills" / "edit-docx"

APP_ID = "merge-test-app"

CONTRACT_PARAS = [
    "SERVICE AGREEMENT",
    "Section 1. Term. This Agreement commences on January 1, 2025 and continues for one year.",
    "Section 2. Payment. Payment shall be due within 30 days of invoice.",
    "Section 3. Termination. Either party may terminate this Agreement with 60 days written notice.",
    "Section 4. Confidentiality. Each party shall keep the other party's information confidential.",
]

AMENDMENT_PARAS = [
    "AMENDMENT TO SERVICE AGREEMENT",
    "This Amendment modifies the Service Agreement as follows.",
    '1. Section 2 (Payment) is amended by replacing "30 days" with "45 days", so that '
    "payment shall be due within 45 days of invoice.",
    "2. Section 3 (Termination) is deleted in its entirety.",
    "3. A new Section 5 (Governing Law) is added: This Agreement shall be governed by the "
    "laws of the State of Delaware.",
]


def _docx_bytes(paragraphs: list[str]) -> bytes:
    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _plain_text(data: bytes) -> str:
    """Visible, non-tracked paragraph text (what survived untouched)."""
    return "\n".join(p.text for p in Document(BytesIO(data)).paragraphs)


def _inserted_text(data: bytes) -> str:
    """Text recorded as tracked insertions (``<w:ins>//<w:t>``)."""
    body = Document(BytesIO(data)).element.body
    return "".join(t.text or "" for ins in body.iter(qn("w:ins")) for t in ins.iter(qn("w:t")))


def _deleted_text(data: bytes) -> str:
    """Text recorded as tracked deletions (``<w:del>//<w:delText>``)."""
    body = Document(BytesIO(data)).element.body
    return "".join(t.text or "" for d in body.iter(qn("w:del")) for t in d.iter(qn("w:delText")))


async def _seed_document(store: LocalFSDocumentStore, doc_id: str, paragraphs: list[str]) -> None:
    """Store a document the way an upload would: extracted text + original .docx bytes.

    ``read_document`` serves the text; ``fetch_document`` serves the bytes under
    ``originals/{doc_id}.docx``.
    """
    await store.save(APP_ID, doc_id, "\n".join(paragraphs))
    await store.save_bytes(APP_ID, f"originals/{doc_id}.docx", _docx_bytes(paragraphs))


@contextlib.contextmanager
def _chdir(path: Path):
    """Run inside *path* so the skill's relative-path scratch files (ops.json,
    merged.docx) stay in the tmp dir instead of littering the repo root."""
    prev = os.getcwd()
    os.chdir(path)
    try:
        yield
    finally:
        os.chdir(prev)


@pytest.mark.asyncio
async def test_merge_amendment_into_contract(tmp_path):
    docstore_root = tmp_path / "docstore"
    store = LocalFSDocumentStore(str(docstore_root))
    await _seed_document(store, "contract", CONTRACT_PARAS)
    await _seed_document(store, "amendment", AMENDMENT_PARAS)

    # Loading from disk installs python-docx into the skill's private venv and wires
    # source_path, so the runner tells the LLM where apply_operations.py lives.
    skill = load_skill_dir(_SKILL_DIR, skill_id="edit-docx")
    assert skill is not None and skill.name == "edit-docx"
    assert skill.site_packages, "python-docx venv was not provisioned for the skill"

    runner = QueryRunner(
        APP_ID,
        _llm,
        RetrievalResources(document_store=store),
        MemoryTiers(),
        skills=[skill],
        max_calls=20,
    )

    query = (
        "Apply the amendment to the original service agreement and produce the merged "
        "Word document. The original contract is doc_id 'contract' and the amendment is "
        "doc_id 'amendment'."
    )

    result = None
    with _chdir(tmp_path):
        async for item in runner.run(query):
            if not isinstance(item, str):
                result = item
    assert result is not None

    # save_artifact writes the redlined output under generated/ in the app's document store.
    artifacts = sorted(glob.glob(str(docstore_root / APP_ID / "generated" / "*.docx")))
    assert artifacts, "no redlined .docx artifact was produced"
    data = Path(artifacts[-1]).read_bytes()
    inserted, deleted, plain = _inserted_text(data), _deleted_text(data), _plain_text(data)

    # replace — payment term redlined: new inserted, old struck
    assert "45 days" in inserted
    assert "30 days" in deleted
    # delete — the termination clause is struck as a tracked deletion
    assert "Either party may terminate" in deleted
    # insert — governing-law clause added as a tracked insertion
    assert "Delaware" in inserted
    # an untouched clause survives as plain, non-tracked text
    assert "Confidentiality" in plain
