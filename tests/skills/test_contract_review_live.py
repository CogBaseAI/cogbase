"""Live end-to-end test for the contract-review skill.

Nothing below the LLM boundary is faked. A real ``LocalFSDocumentStore`` backs the
document tools, the skill is loaded from disk with ``load_skill_dir`` — which
installs its declared ``python-docx`` dependency into the skill's private venv —
and the ``QueryRunner`` drives the real agent loop: routing to the skill, reading
the contract, fetching the original ``.docx``, shelling out to the bundled
``segment_clauses.py`` and ``build_ops.py``, and saving the review file.

The test builds a small multi-clause contract as a real ``.docx``, seeds it the
way an upload would (extracted text + original bytes), and drives the two phases
of the skill in two conversation turns:

  1. **Analyze** — ask the agent to review the contract **for the Client from a
     disadvantaged position** (both supplied up front so the skill's confirmation
     step is satisfied without an interactive turn). Re-open the produced review
     ``ops.json`` and assert the lens is recorded, clauses carry well-formed risk
     levels and at least one suggested change, and **every anchor the LLM produced
     round-trips through edit-docx with zero unmatched**.

  2. **Apply + refine (Phase 2)** — continue the conversation and tell the agent to
     accept one clause's suggestion, reject the rest, and produce the redline. Then
     assert the agent drove ``build_ops.py patch`` + ``to-edit-ops`` +
     ``edit-docx`` through the loop to produce a real tracked-changes ``.docx`` that
     reflects **only the accepted verdict** — the end-to-end proof that the
     verdict-gated apply path works on real model output.

Run with::

    pytest tests/skills/test_contract_review_live.py -m live
"""

from __future__ import annotations

import contextlib
import glob
import importlib.util
import json
import os
from io import BytesIO
from pathlib import Path
from xml.sax.saxutils import escape

import pytest

pytest.importorskip("docx")  # needed to build the fixture and re-apply the anchors

from docx import Document  # noqa: E402  (import after importorskip)

from cogbase.core.query_runner import (  # noqa: E402
    MemoryTiers,
    QueryResult,
    QueryRunner,
    RetrievalResources,
)
from cogbase.skills.skill import load_skill_dir  # noqa: E402
from cogbase.stores.document.local_fs import LocalFSDocumentStore  # noqa: E402
from tests.live_setup import make_llm  # noqa: E402

_llm = make_llm()

pytestmark = [
    pytest.mark.live,
    pytest.mark.skipif(_llm is None, reason="No LLM configured (see tests/live_setup)"),
]

_SKILLS = Path(__file__).resolve().parents[2] / "skills"
_SKILL_DIR = _SKILLS / "contract-review"

APP_ID = "contract-review-test-app"
DOC_ID = "msa"


def _load(name: str, path: Path):
    spec = importlib.util.spec_from_file_location(name, path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


build_mod = _load("build_ops", _SKILL_DIR / "build_ops.py")
edit_helper = _load("apply_operations", _SKILLS / "edit-docx" / "apply_operations.py")

CONTRACT_PARAS = [
    "MASTER SERVICES AGREEMENT",
    "This Master Services Agreement is entered into between Acme Provider Inc. "
    "(the \"Provider\") and Beta Client LLC (the \"Client\").",
    "1. PAYMENT TERMS",
    "Payment shall be due within 15 days of invoice receipt by the Client, and late "
    "payments accrue interest at 5% per month.",
    "2. LIABILITY",
    "The Client's total liability under this Agreement is unlimited, while the Provider's "
    "liability is capped at the fees paid in the prior month.",
    "3. TERMINATION",
    "The Provider may terminate this Agreement at any time for convenience with 5 days "
    "notice; the Client may only terminate for material breach.",
    "4. INDEMNIFICATION",
    "The Client shall indemnify the Provider against any and all claims arising from the "
    "use of the services, without limitation.",
    "5. GOVERNING LAW",
    "This Agreement is governed by the laws of the State of New York.",
]


def _docx_bytes(paragraphs: list[str]) -> bytes:
    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def _seed_document(store: LocalFSDocumentStore, doc_id: str, paragraphs: list[str]) -> None:
    """Store a document the way an upload would: extracted text + original .docx bytes."""
    await store.save(APP_ID, doc_id, "\n".join(paragraphs))
    await store.save_bytes(APP_ID, f"originals/{doc_id}.docx", _docx_bytes(paragraphs))


@contextlib.contextmanager
def _chdir(path: Path):
    """Run inside *path* so the skill's relative scratch files (clauses.json,
    analysis.json, review.json) stay in the tmp dir, not the repo root."""
    prev = os.getcwd()
    os.chdir(path)
    try:
        yield
    finally:
        os.chdir(prev)


async def _run_turn(runner: QueryRunner, query: str, history: list[dict]) -> QueryResult:
    """Drive one agent turn to completion, returning its final QueryResult.

    The runner streams ``str`` tokens then yields a final ``QueryResult``; memory
    is off in this test, so we thread continuity by passing prior turns as
    ``history``.
    """
    result: QueryResult | None = None
    async for item in runner.run(query, history=history):
        if not isinstance(item, str):
            result = item
    assert result is not None, "runner produced no final QueryResult"
    return result


def _generated(docstore_root: Path, pattern: str) -> list[str]:
    return sorted(glob.glob(str(docstore_root / APP_ID / "generated" / pattern)))


def _document_xml(path: str) -> str:
    """Return the raw word/document.xml of a .docx (so tracked-change runs are visible).

    python-docx's ``paragraph.runs`` skips runs nested inside ``<w:ins>``/``<w:del>``,
    so inspecting the redline's insertions/deletions means reading the part directly.
    """
    import zipfile

    with zipfile.ZipFile(path) as z:
        return z.read("word/document.xml").decode("utf-8")


@pytest.mark.asyncio
async def test_review_contract_produces_applicable_review(tmp_path):
    docstore_root = tmp_path / "docstore"
    store = LocalFSDocumentStore(str(docstore_root))
    await _seed_document(store, DOC_ID, CONTRACT_PARAS)

    # Loading from disk installs python-docx into the skill's private venv and wires
    # source_path, so the runner tells the LLM where the helper scripts live.
    skill = load_skill_dir(_SKILL_DIR, skill_id="contract-review")
    assert skill is not None and skill.name == "contract-review"
    assert skill.site_packages, "python-docx venv was not provisioned for the skill"

    runner = QueryRunner(
        APP_ID,
        _llm,
        RetrievalResources(document_store=store),
        MemoryTiers(),
        skills=[skill],
        max_calls=30,
    )

    # --- Turn 1: analyze ----------------------------------------------------
    review_query = (
        f"Review the contract with doc_id '{DOC_ID}'. It represents the Client "
        "(Beta Client LLC), and the review position is disadvantaged. Both the "
        "represented party and the position are confirmed — proceed with the full "
        "clause-by-clause review and save the review file without asking me to "
        "confirm again."
    )

    with _chdir(tmp_path):
        review_result = await _run_turn(runner, review_query, history=[])

    # The skill persists the review as a JSON artifact under generated/.
    review_files = _generated(docstore_root, "*.json")
    assert review_files, "no review ops.json artifact was produced"
    review_id = os.path.basename(review_files[-1])  # the artifact id is the filename
    review = json.loads(Path(review_files[-1]).read_text())

    # Review lens recorded as requested.
    meta = review["meta"]
    assert "client" in (meta.get("representative_party") or "").lower()
    assert meta.get("review_position") == "disadvantaged"

    # Well-formed clauses with valid risk levels.
    clauses = review["clauses"]
    assert clauses, "review must contain clauses"
    assert all(c["risk"]["level"] in {"high", "medium", "low", "none"} for c in clauses)

    # The disadvantaged Client has plenty to push back on — expect suggested changes.
    suggested = [c for c in clauses if c.get("suggestion")]
    assert suggested, "expected at least one suggested change for a disadvantaged party"

    # End-to-end anchoring proof: every anchor the LLM produced must locate its
    # paragraph in the base, so the whole review is applicable as a redline.
    ops = build_mod.to_edit_ops(review, accepted_only=False)["operations"]
    assert ops
    doc = Document(BytesIO(_docx_bytes(CONTRACT_PARAS)))
    report = edit_helper.apply_operations(doc, ops, "test", "2026-01-01T00:00:00Z")
    unmatched = [r for r in report if not r["matched"]]
    assert not unmatched, f"LLM produced anchors that don't match the base: {unmatched}"

    # --- Turn 2: apply + refine (Phase 2) -----------------------------------
    # Accept exactly one suggestion and reject the rest, so the redline must be
    # verdict-gated: the accepted change appears, the rejected ones do not.
    accept_clause = suggested[0]
    reject_clauses = suggested[1:]
    accept_id = accept_clause["clause_id"]
    reject_ids = [c["clause_id"] for c in reject_clauses]
    accept_new_text = accept_clause["suggestion"]["new_text"]

    reject_instruction = (
        f"reject the suggestions for clauses {', '.join(reject_ids)}"
        if reject_ids
        else "leave the other clauses as they are"
    )
    apply_query = (
        f"Finalize the review (its artifact id is '{review_id}'). Accept the suggestion "
        f"for clause {accept_id} and {reject_instruction}. Then apply only the accepted "
        "changes and give me the tracked-changes redline .docx — accepted verdicts only, "
        "not a preview of everything."
    )
    history = [
        {"role": "user", "content": review_query},
        {"role": "assistant", "content": review_result.answer},
    ]

    with _chdir(tmp_path):
        await _run_turn(runner, apply_query, history=history)

    # A redline .docx artifact was produced.
    redlines = _generated(docstore_root, "*.docx")
    assert redlines, "no redline .docx artifact was produced by the apply turn"
    redline_path = redlines[-1]
    Document(redline_path)  # opens cleanly — a valid .docx
    xml = _document_xml(redline_path)

    # It is a tracked-changes redline carrying the accepted change...
    assert "<w:ins" in xml, "redline has no tracked insertions"
    assert escape(accept_new_text) in xml, "accepted clause's new text is not in the redline"

    # ...and the verdict gate held: rejected suggestions were not applied.
    for clause in reject_clauses:
        rejected_text = clause["suggestion"].get("new_text")
        if rejected_text:
            assert escape(rejected_text) not in xml, (
                f"rejected clause {clause['clause_id']} leaked into the redline"
            )

    # The patched review persisted the verdicts deterministically. Artifact ids
    # carry a random suffix, so file order isn't chronological — take each clause's
    # *decided* verdict across all review artifacts (a later patch only ever moves a
    # clause off ``pending``, never back onto it).
    verdict_by_id: dict[str, str] = {}
    for path in _generated(docstore_root, "*.json"):
        for clause in json.loads(Path(path).read_text()).get("clauses", []):
            verdict = clause.get("verdict")
            if verdict and verdict != "pending":
                verdict_by_id[clause["clause_id"]] = verdict
    assert verdict_by_id.get(accept_id) == "accepted"
    for reject_id in reject_ids:
        assert verdict_by_id.get(reject_id) == "rejected"
